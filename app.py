from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import os

# ---------------- CONFIG ----------------
COL1_WIDTH = 0.4
COL2_WIDTH = 1.8
COL3_WIDTH = 5.2

NORMAL_FONT = 10
HEADER_FONT = 11
BIG_HEADER_FONT = 12
# --------------------------------------


def add_hyperlink(paragraph, url, text, bold=False):
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_big_space(p):
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)


def add_small_space(p):
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)


# ---------------- DOCUMENT ----------------
os.makedirs("output", exist_ok=True)
doc = Document()

# PAGE SETUP (MUST BE BEFORE TABLE CREATION)
section = doc.sections[0]
section.page_width = Inches(8.27)
section.page_height = Inches(11.69)
section.top_margin = Inches(0.6)
section.bottom_margin = Inches(0.6)
section.left_margin = Inches(0.5)
section.right_margin = Inches(0.5)

# ---------------- HEADER ----------------
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def header(text, size, bold=False):
    r = p.add_run(text + "\n")
    r.font.size = Pt(size)
    r.bold = bold

header("FORM ‘A’", BIG_HEADER_FONT, True)
header("MEDIATION APPLICATION FORM", BIG_HEADER_FONT, True)
header("[REFER RULE 3(1)]", HEADER_FONT, True)
header("Mumbai District Legal Services Authority", NORMAL_FONT)
header("City Civil Court, Mumbai", NORMAL_FONT)

# ---------------- TABLE ----------------
table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
table.autofit = False

table._tbl.remove(table.rows[0]._tr)

table.columns[0].width = Inches(COL1_WIDTH)
table.columns[1].width = Inches(COL2_WIDTH)
table.columns[2].width = Inches(COL3_WIDTH)


def add_row(c1="", c2="", c3="", bold=False, merge=False):
    row = table.add_row().cells
    row[0].text = c1
    row[1].text = c2
    row[2].text = c3

    if merge:
        row[0].merge(row[1]).merge(row[2])

    for cell in row:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(NORMAL_FONT)
                r.bold = bold
            add_small_space(p)


# ---------------- DETAILS OF PARTIES ----------------
row = table.add_row().cells
row[0].merge(row[1]).merge(row[2])
p = row[0].paragraphs[0]
p.add_run("DETAILS OF PARTIES:").bold = True
add_big_space(p)

add_row("1", "Name of Applicant", "{{client_name}}", bold=True)

# Address heading
row = table.add_row().cells
row[0].text = ""

p = row[1].paragraphs[0]
p.add_run("Address and contact details of Applicant").bold = True
add_big_space(p)  # <-- THIS adds height

row[1].merge(row[2])


# Address content
row = table.add_row().cells
row[0].paragraphs[0].add_run("1").bold = True
row[1].paragraphs[0].add_run("Address").bold = True

p = row[2].paragraphs[0]
add_big_space(p)

p.add_run("REGISTERED ADDRESS:\n").bold = True
p.add_run("{{branch_address}}\n\n")
p.add_run("CORRESPONDENCE BRANCH ADDRESS:\n").bold = True
p.add_run("{{branch_address}}")

add_small_space(p)

add_row("", "Telephone No.", "{{mobile}}", bold=True)
add_row("", "Mobile No.", "", bold=True)

# Email
row = table.add_row().cells
row[0].text = ""
row[1].paragraphs[0].add_run("Email ID").bold = True
p = row[2].paragraphs[0]
add_hyperlink(p, "mailto:info@kslegal.co.in", "info@kslegal.co.in", bold=True)
add_small_space(p)

# ---------------- OPPOSITE PARTY ----------------
row = table.add_row().cells
row[0].text = "2"
p = row[1].paragraphs[0]
p.add_run("Name, Address and Contact details of Opposite Party:").bold = True
add_big_space(p)  
row[1].merge(row[2])

row = table.add_row().cells
row[0].text = ""
p = row[1].paragraphs[0]
p.add_run("Address and contact details of Defendant/s").bold = True
add_big_space(p)  
row[1].merge(row[2])

add_row("", "Name", "{{customer_name}}", bold=True)

row = table.add_row().cells
row[0].text = ""
row[1].paragraphs[0].add_run("Address").bold = True

p = row[2].paragraphs[0]
p.add_run("REGISTERED ADDRESS:\n").bold = True
p.add_run("{% if address1 and address1 != '' %}{{address1}}{% else %}________________{% endif %}\n\n")
p.add_run("CORRESPONDENCE ADDRESS:\n").bold = True
p.add_run("{% if address1 and address1 != '' %}{{address1}}{% else %}________________{% endif %}")

add_row("", "Telephone No.", "", bold=True)
add_big_space(p)
add_row("", "Mobile No.", "", bold=True)
add_big_space(p)
add_row("", "Email ID", "", bold=True)
add_big_space(p)

# ---------------- DISPUTE ----------------
row = table.add_row().cells
row[0].merge(row[1]).merge(row[2])
p = row[0].paragraphs[0]
p.add_run("DETAILS OF DISPUTE:").bold = True
add_big_space(p)

row = table.add_row().cells
row[0].merge(row[1]).merge(row[2])
p = row[0].paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("THE COMM. COURTS (PRE-INSTITUTION SETTLEMENT) RULES, 2018")
run.bold = True
run.underline = True
add_big_space(p)

row = table.add_row().cells
row[0].text = ""
row[1].text = "Nature of disputes as per section 2(1)(c) of the Commercial Courts Act, 2015 (4 of 2016):"
row[1].merge(row[2])
p = row[1].paragraphs[0]
run = p.runs[0]
run.bold = True
run.font.size = Pt(10)

# ---------------- SAVE ----------------
doc.save("output/Mediation_Form.docx")
print("✅ Document created successfully!")

